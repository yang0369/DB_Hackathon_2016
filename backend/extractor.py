import io
import os
import pypdf
import pdfplumber
import docx

def extract_text_from_pdf(file_bytes: bytes) -> str:
    """Extract text from PDF file bytes using pdfplumber with pypdf fallback."""
    text_content = []
    
    # Try pdfplumber first for better layout preservation
    try:
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text_content.append(page_text)
        if text_content:
            return "\n\n".join(text_content)
    except Exception as e:
        print(f"pdfplumber extraction failed/fallback to pypdf: {e}")

    # Fallback to pypdf
    try:
        reader = pypdf.PdfReader(io.BytesIO(file_bytes))
        for page in reader.pages:
            extracted = page.extract_text()
            if extracted:
                text_content.append(extracted)
        return "\n\n".join(text_content)
    except Exception as e:
        print(f"pypdf extraction failed: {e}")
        return ""

def extract_text_from_docx(file_bytes: bytes) -> str:
    """Extract text from Microsoft Word DOCX file bytes."""
    try:
        doc = docx.Document(io.BytesIO(file_bytes))
        full_text = []
        for para in doc.paragraphs:
            if para.text:
                full_text.append(para.text)
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join([cell.text.strip() for cell in row.cells if cell.text.strip()])
                if row_text:
                    full_text.append(row_text)
        return "\n".join(full_text)
    except Exception as e:
        print(f"DOCX extraction failed: {e}")
        return ""

def extract_text_from_file(file_bytes: bytes, filename: str) -> str:
    """Extract text content based on file extension."""
    ext = os.path.splitext(filename)[1].lower()
    if ext == ".pdf":
        return extract_text_from_pdf(file_bytes)
    elif ext in [".docx", ".doc"]:
        return extract_text_from_docx(file_bytes)
    elif ext in [".txt", ".json", ".md"]:
        return file_bytes.decode("utf-8", errors="ignore")
    else:
        # Fallback to UTF-8 text decoding
        return file_bytes.decode("utf-8", errors="ignore")
